"""Tests for common.document_reader module."""

from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch


class TestLibraryAvailability:
    """Tests for library availability checks."""

    def test_is_docx_available_returns_true(self):
        """Should return True when python-docx is installed."""
        from deriva.common.document_reader import is_docx_available

        assert is_docx_available() is True

    def test_is_pdf_available_returns_true(self):
        """Should return True when pypdf is installed."""
        from deriva.common.document_reader import is_pdf_available

        assert is_pdf_available() is True


class TestReadDocument:
    """Tests for read_document routing function."""

    def test_routes_docx_to_read_docx(self):
        """Should route .docx files to read_docx."""
        from deriva.common.document_reader import read_document

        with patch("deriva.common.document_reader.read_docx") as mock_read:
            mock_read.return_value = "docx content"
            result = read_document(Path("test.docx"))
            mock_read.assert_called_once()
            assert result == "docx content"

    def test_routes_pdf_to_read_pdf(self):
        """Should route .pdf files to read_pdf."""
        from deriva.common.document_reader import read_document

        with patch("deriva.common.document_reader.read_pdf") as mock_read:
            mock_read.return_value = "pdf content"
            result = read_document(Path("test.pdf"))
            mock_read.assert_called_once()
            assert result == "pdf content"

    def test_returns_none_for_unknown_extension(self):
        """Should return None for unsupported file types."""
        from deriva.common.document_reader import read_document

        result = read_document(Path("test.txt"))
        assert result is None

    def test_handles_uppercase_extensions(self):
        """Should handle uppercase extensions."""
        from deriva.common.document_reader import read_document

        with patch("deriva.common.document_reader.read_docx") as mock_read:
            mock_read.return_value = "content"
            result = read_document(Path("test.DOCX"))
            mock_read.assert_called_once()
            assert result == "content"


class TestReadDocx:
    """Tests for read_docx function."""

    def test_returns_none_when_library_unavailable(self):
        """Should return None when python-docx is not available."""
        from deriva.common.document_reader import read_docx

        with patch("deriva.common.document_reader.is_docx_available", return_value=False):
            result = read_docx(Path("test.docx"))
            assert result is None

    def test_returns_none_on_read_error(self):
        """Should return None when reading fails."""
        from deriva.common.document_reader import read_docx

        with patch("deriva.common.document_reader.is_docx_available", return_value=True):
            # Try to read a non-existent file
            result = read_docx(Path("nonexistent.docx"))
            assert result is None

    def test_extracts_paragraphs(self):
        """Should extract paragraph text."""
        import docx

        from deriva.common.document_reader import read_docx

        # Create a temporary directory to avoid file locking issues
        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("Second paragraph")
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "First paragraph" in result
            assert "Second paragraph" in result

    def test_extracts_headings(self):
        """Should extract heading styles as markdown headings."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            doc.add_heading("Title", level=1)
            doc.add_heading("Section", level=2)
            doc.add_paragraph("Content")
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "# Title" in result
            assert "## Section" in result
            assert "Content" in result

    def test_extracts_title_style(self):
        """Should extract Title style as h1."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            title = doc.add_paragraph("Document Title")
            title.style = "Title"
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "# Document Title" in result

    def test_extracts_list_items(self):
        """Should extract list items with markdown bullets."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            item = doc.add_paragraph("Item one")
            item.style = "List Bullet"
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "- Item one" in result

    def test_extracts_deeper_heading_levels(self):
        """Should extract heading levels 3-6 correctly."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            doc.add_heading("Level 3", level=3)
            doc.add_heading("Level 4", level=4)
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "### Level 3" in result
            assert "#### Level 4" in result

    def test_extracts_bold_text(self):
        """Should format bold text with markdown."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            para = doc.add_paragraph()
            run = para.add_run("Bold text")
            run.bold = True
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "**Bold text**" in result

    def test_extracts_italic_text(self):
        """Should format italic text with markdown."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            para = doc.add_paragraph()
            run = para.add_run("Italic text")
            run.italic = True
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "*Italic text*" in result

    def test_extracts_bold_italic_text(self):
        """Should format bold+italic text with markdown."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            para = doc.add_paragraph()
            run = para.add_run("Bold italic")
            run.bold = True
            run.italic = True
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "***Bold italic***" in result

    def test_extracts_tables(self):
        """Should extract tables as markdown tables."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Header1"
            table.cell(0, 1).text = "Header2"
            table.cell(1, 0).text = "Data1"
            table.cell(1, 1).text = "Data2"
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "Header1" in result
            assert "Header2" in result
            assert "Data1" in result
            assert "---" in result  # Table separator

    def test_skips_empty_paragraphs(self):
        """Should skip empty paragraphs."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            doc.add_paragraph("")  # Empty
            doc.add_paragraph("Content")
            doc.add_paragraph("   ")  # Whitespace only
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            # Should have content but not excessive blank lines
            assert "Content" in result


class TestReadPdf:
    """Tests for read_pdf function."""

    def test_returns_none_when_library_unavailable(self):
        """Should return None when pypdf is not available."""
        from deriva.common.document_reader import read_pdf

        with patch("deriva.common.document_reader.is_pdf_available", return_value=False):
            result = read_pdf(Path("test.pdf"))
            assert result is None

    def test_returns_none_on_read_error(self):
        """Should return None when reading fails."""
        from deriva.common.document_reader import read_pdf

        with patch("deriva.common.document_reader.is_pdf_available", return_value=True):
            result = read_pdf(Path("nonexistent.pdf"))
            assert result is None

    def test_extracts_text_from_pages_with_mock(self):
        """Should extract text with page headers."""
        from deriva.common.document_reader import read_pdf

        # Mock a PDF reader
        mock_reader = MagicMock()
        mock_page1 = MagicMock()
        mock_page1.extract_text.return_value = "Page 1 content"
        mock_page2 = MagicMock()
        mock_page2.extract_text.return_value = "Page 2 content"
        mock_reader.pages = [mock_page1, mock_page2]

        with (
            patch("deriva.common.document_reader.is_pdf_available", return_value=True),
            patch("pypdf.PdfReader", return_value=mock_reader),
        ):
            result = read_pdf(Path("test.pdf"))

            assert result is not None
            assert "## Page 1" in result
            assert "Page 1 content" in result
            assert "## Page 2" in result
            assert "Page 2 content" in result

    def test_skips_empty_pages(self):
        """Should skip pages with no text."""
        from deriva.common.document_reader import read_pdf

        # Mock a PDF reader with an empty page
        mock_reader = MagicMock()
        mock_page_empty = MagicMock()
        mock_page_empty.extract_text.return_value = ""
        mock_page_content = MagicMock()
        mock_page_content.extract_text.return_value = "Some content"
        mock_reader.pages = [mock_page_empty, mock_page_content]

        with (
            patch("deriva.common.document_reader.is_pdf_available", return_value=True),
            patch("pypdf.PdfReader", return_value=mock_reader),
        ):
            result = read_pdf(Path("test.pdf"))

            assert result is not None
            assert "## Page 2" in result  # Second page has content
            assert "## Page 1" not in result  # First page is empty

    def test_skips_whitespace_only_pages(self):
        """Should skip pages with only whitespace."""
        from deriva.common.document_reader import read_pdf

        mock_reader = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "   \n\n   "  # Whitespace only
        mock_reader.pages = [mock_page]

        with (
            patch("deriva.common.document_reader.is_pdf_available", return_value=True),
            patch("pypdf.PdfReader", return_value=mock_reader),
        ):
            result = read_pdf(Path("test.pdf"))

            # Should return empty string since all pages are whitespace
            assert result == ""


class TestTableToMarkdown:
    """Tests for _table_to_markdown helper function."""

    def test_escapes_pipe_characters(self):
        """Should escape pipe characters in cell content."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Value|with|pipes"
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            # Pipes should be escaped
            assert "\\|" in result


class TestFormatParagraphRuns:
    """Tests for _format_paragraph_runs helper function."""

    def test_handles_empty_runs(self):
        """Should handle runs with empty text."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            para = doc.add_paragraph()
            para.add_run("")  # Empty run
            para.add_run("Content")
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "Content" in result

    def test_combines_multiple_runs(self):
        """Should combine multiple runs in a paragraph."""
        import docx

        from deriva.common.document_reader import read_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            filepath = Path(tmpdir) / "test.docx"
            doc = docx.Document()
            para = doc.add_paragraph()
            para.add_run("Normal ")
            bold = para.add_run("bold")
            bold.bold = True
            para.add_run(" more")
            doc.save(str(filepath))

            result = read_docx(filepath)

            assert result is not None
            assert "Normal " in result
            assert "**bold**" in result
            assert " more" in result
